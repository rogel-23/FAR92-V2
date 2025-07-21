import streamlit as st
import pandas as pd
import os
from datetime import datetime
import json
from io import BytesIO
import math
from datetime import datetime, date
from google_drive_utils import upload_to_drive
from streamlit import secrets
import uuid
st.write("Clés dans st.secrets :", list(st.secrets.keys()))
from supabase import create_client
from google_drive_utils import list_rapports_for_arbitre  


def upload_rapport_to_supabase(uploaded_file, arbitre_id):
    """
    Upload un fichier PDF dans Supabase dans un dossier nommé par l'ID de l'arbitre.
    Renvoie l'URL publique du fichier.
    """
    SUPABASE_URL = st.secrets["SUPABASE_URL"]
    SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
    supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
    bucket = "rapports"

    filepath = f"{arbitre_id}/{uploaded_file.name}"

    # Supprimer l’ancien fichier s’il existe
    delete_res = supabase.storage.from_(bucket).remove([filepath])
    if getattr(delete_res, "error", None):
        if delete_res.error.message != "The resource does not exist":
            raise Exception(f"Erreur lors de la suppression : {delete_res.error.message}")

    # Upload avec écrasement autorisé
    try:
        supabase.storage.from_(bucket).remove([filepath])
    except:
        pass

    res = supabase.storage.from_(bucket).upload(filepath, uploaded_file.getvalue())
    if hasattr(res, "status_code") and res.status_code == 409:
        raise Exception("Un fichier du même nom existe déjà. Veuillez le renommer ou le supprimer avant l'upload.")
    elif hasattr(res, "get") and res.get("error"):
        raise Exception(f"Erreur Supabase : {res['error']['message']}")



    public_url = supabase.storage.from_(bucket).get_public_url(filepath)
    return public_url

def list_rapports_for_arbitre(arbitre_id):
    """
    Retourne la liste des fichiers (URL publiques) déjà associés à un arbitre.
    """
    SUPABASE_URL = st.secrets["SUPABASE_URL"]
    SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
    supabase = create_client(SUPABASE_URL, SUPABASE_KEY)
    bucket = "rapports"

    # Liste des fichiers dans le dossier de l’arbitre
    res = supabase.storage.from_(bucket).list(path=arbitre_id)

    if getattr(res, "error", None):
        raise Exception(f"Erreur Supabase (list) : {res.error.message}")

    fichiers = res  # Liste d’objets {'name': ..., 'created_at': ..., ...}
    
    # Génère les URLs publiques
    urls = []
    for fichier in fichiers:
        path = f"{arbitre_id}/{fichier['name']}"
        url = supabase.storage.from_(bucket).get_public_url(path)
        urls.append((fichier["name"], url))
    
    return urls

folder_id = "1Oe6hhlJuU2S8cK_u-eo-tdig1t8onhl_"  # fallback

# === Initialisation session_state ===
if "far_arbitres" not in st.session_state:
    st.session_state["far_arbitres"] = []

# === FONCTIONS DE CHARGEMENT/SAUVEGARDE ===
APP_DIR = os.path.dirname(os.path.abspath(__file__))
FILENAME = os.path.join(APP_DIR, "far_arbitres.xlsx")
os.makedirs("rapports", exist_ok=True)

def load_arbitres():
    if os.path.exists(FILENAME):
        df = pd.read_excel(FILENAME)
        for col in ["Rassemblements"]:
            if col not in df.columns:
                df[col] = ""
        return df.to_dict(orient="records")
    else:
        df = pd.DataFrame(columns=[
            "Nom", "Prénom", "Catégorie", "Date de naissance", "Âge",
            "Club", "Téléphone", "Email", "Rassemblements"
        ])
        df.to_excel(FILENAME, index=False)
        return []


def save_arbitres(data):
    df = pd.DataFrame(data)
    df.to_excel(FILENAME, index=False)

# === TITRE + BOUTON SAUVEGARDE ALIGNÉS ===
col_title, col_save = st.columns([5, 1])

with col_title:
    st.title("⚽ FAR 92 - Application de gestion")
    st.markdown("Bienvenue sur l'application officielle de la **Filière Arbitrage Régionale du District 92**.")

with col_save:
    st.write("")  # décalage vertical
    df = pd.DataFrame(st.session_state["far_arbitres"])

    # Test : exécution locale ou non
    is_local = os.path.exists(".git") or os.getenv("STREAMLIT_ENV") != "cloud"

    if is_local:
        if st.button("💾 Sauvegarder"):
            df.to_excel("far_arbitres.xlsx", index=False)
            st.success("Fichier Excel sauvegardé dans le dossier du projet.")
    else:
        buffer = BytesIO()
        df.to_excel(buffer, index=False)
        buffer.seek(0)
        st.download_button(
            "📥 Télécharger Excel",
            data=buffer,
            file_name="far_arbitres.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )


# === CHARGEMENT INITIAL D'UN FICHIER EXCEL ===
if "fichier_source" not in st.session_state:
    st.session_state["fichier_source"] = None

if st.session_state["fichier_source"] is None:
    st.subheader("📂 Charger un fichier Excel FAR")
    uploaded_file = st.file_uploader("Sélectionnez un fichier Excel", type=["xlsx"])

    if uploaded_file is not None:
        df = pd.read_excel(uploaded_file)
        for col in ["Rassemblements"]:
            if col not in df.columns:
                df[col] = ""
        st.session_state["far_arbitres"] = df.to_dict(orient="records")
        st.session_state["fichier_source"] = uploaded_file.name
        st.success("Fichier chargé avec succès.")
        st.rerun()
    else:
        st.stop()


# === SESSION STATE ===
if "far_arbitres" not in st.session_state:
    st.session_state["far_arbitres"] = load_arbitres()

# === MENU ===
st.subheader("📚 Modules disponibles")
col1, col2, col3 = st.columns(3)

with col1:
    action = st.radio("Menu", [
    "📝 Compte-rendu rassemblement",
    "📊 Récapitulatif des rassemblements",
    "📝 Saisie des examens",
    "📊 Récapitulatif des examens",
    "🛑 Ajouter des manquements",
    "📉 Récapitulatif des manquements",
    "📎 Déposer un rapport d'observation",
    "👤 Fiche arbitre",
    "➕ Ajouter / ❌ Supprimer un arbitre"
])

# === AJOUT / SUPPRESSION ===
if action == "➕ Ajouter / ❌ Supprimer un arbitre":
    st.subheader("➕ Ajouter un arbitre")
    with st.form("ajout_arbitre_form"):
        nom = st.text_input("Nom").upper()
        prenom = st.text_input("Prénom")
        categorie = st.selectbox("Catégorie FAR", ["FAR-S1", "FAR-S2", "FAR-A1", "FAR-J1", "FAR-F1"])
        date_naissance = st.date_input("Date de naissance", min_value=date(1900, 1, 1))
        club = st.text_input("Club")
        tel = st.text_input("Téléphone")
        email = st.text_input("Adresse mail")
        submit = st.form_submit_button("Ajouter")

        if submit and nom and prenom:
            today = datetime.date.today()
            age = today.year - date_naissance.year - ((today.month, today.day) < (date_naissance.month, date_naissance.day))
            st.session_state["far_arbitres"].append({
                "Nom": nom, "Prénom": prenom, "Catégorie": categorie,
                "Date de naissance": date_naissance.strftime("%d/%m/%Y"), "Âge": age,
                "Club": club, "Téléphone": tel, "Email": email,
                "Rassemblements": ""
            })
            save_arbitres(st.session_state["far_arbitres"])
            st.success("Ajouté avec succès.")

    st.subheader("❌ Supprimer un arbitre")
    for i, a in enumerate(st.session_state["far_arbitres"]):
        col1, col2, col3, col4 = st.columns([3, 3, 3, 1])
        col1.markdown(f"**{a['Prénom']} {a['Nom']}**")
        col2.write(a["Catégorie"])
        col3.write(a["Club"])
        if col4.button("🗑️", key=f"del_{i}"):
            st.session_state["far_arbitres"].pop(i)
            save_arbitres(st.session_state["far_arbitres"])
            st.rerun()

# === EXPORT ===
elif action == "📄 Exporter la liste modifiée":
    st.subheader("📄 Export")
    df = pd.DataFrame(st.session_state["far_arbitres"])
    buffer = BytesIO()
    df.to_excel(buffer, index=False)
    buffer.seek(0)
    st.download_button("📥 Télécharger le fichier", data=buffer, file_name="far_arbitres.xlsx")

# === COMPTE-RENDU RASSEMBLEMENT ===
elif action == "📝 Compte-rendu rassemblement":
    st.subheader("📝 Nouveau compte-rendu")
    type_rass = st.selectbox("Type de rassemblement", ["Réunion", "Stage", "Test physique", "Autre"])

    if type_rass in ["Réunion", "Autre"]:
        with st.form("form_reunion"):
            nom_rass = st.text_input("Nom de la réunion")
            date_rass = st.date_input("Date de la réunion")
            statuts = {}
            st.markdown("### Présence des arbitres")
            commentaires = {}
            for i, a in enumerate(st.session_state["far_arbitres"]):
                nom_complet = f"{a['Prénom']} {a['Nom']}"
                st.markdown(f"**{nom_complet}**")
                col1, col2 = st.columns([3, 3])
                statut = col1.selectbox("Statut", ["Présent", "Absent excusé", "Absent non excusé"], key=f"r_statut_{i}")
                commentaire = col2.text_input("Commentaire individuel", key=f"r_comment_{i}")
                statuts[nom_complet] = statut
                commentaires[nom_complet] = commentaire.strip()

            observations = st.text_area("Observations générales (facultatif)")
            submit = st.form_submit_button("Enregistrer")

            if submit and nom_rass:
                for i, a in enumerate(st.session_state["far_arbitres"]):
                    rass = json.loads(a.get("Rassemblements", "") or "[]")
                    rass = [r for r in rass if r.get("Nom") != nom_rass]

                    rass.append({
                        "Nom": nom_rass,
                        "Type": type_rass,
                        "Date": date_rass.strftime("%d/%m/%Y"),
                        "Statut": statuts[f"{a['Prénom']} {a['Nom']}"],
                        "Observations": observations,
                        "Observations individuelles": commentaires.get(f"{a['Prénom']} {a['Nom']}", "")
                    })
                    st.session_state["far_arbitres"][i]["Rassemblements"] = json.dumps(rass)
                save_arbitres(st.session_state["far_arbitres"])
                st.success("Réunion enregistrée avec succès.")

    elif type_rass == "Stage":
        with st.form("form_stage"):
            nom_stage = st.text_input("Nom du stage")
            date_debut = st.date_input("Date de début")
            date_fin = st.date_input("Date de fin")
            statuts = {}
            st.markdown("### Présence des arbitres")
            commentaires = {}
            for i, a in enumerate(st.session_state["far_arbitres"]):
                nom_complet = f"{a['Prénom']} {a['Nom']}"
                st.markdown(f"**{nom_complet}**")
                col1, col2 = st.columns([3, 3])
                statut = col1.selectbox("Statut", ["Présent", "Absent excusé", "Absent non excusé"], key=f"s_statut_{i}")
                commentaire = col2.text_input("Commentaire individuel", key=f"s_comment_{i}")
                statuts[nom_complet] = statut
                commentaires[nom_complet] = commentaire.strip()


            observations = st.text_area("Observations générales (facultatif)")
            submit = st.form_submit_button("Enregistrer")

            if submit and nom_stage:
                for i, a in enumerate(st.session_state["far_arbitres"]):
                    rass = json.loads(a.get("Rassemblements", "") or "[]")
                    rass = [r for r in rass if r.get("Nom") != nom_stage]

                    rass.append({
                        "Nom": nom_stage,
                        "Type": "Stage",
                        "Date début": date_debut.strftime("%d/%m/%Y"),
                        "Date fin": date_fin.strftime("%d/%m/%Y"),
                        "Statut": statuts[f"{a['Prénom']} {a['Nom']}"],
                        "Observations": observations,
                        "Observations individuelles": commentaires.get(f"{a['Prénom']} {a['Nom']}", "")

                    })
                    st.session_state["far_arbitres"][i]["Rassemblements"] = json.dumps(rass)
                save_arbitres(st.session_state["far_arbitres"])
                st.success("Stage enregistré avec succès.")

    elif type_rass == "Test physique":
        commentaires = {}
        with st.form("form_test_physique"):
            nom_test = st.text_input("Nom du test physique")
            date_test = st.date_input("Date du test")
            statuts = {}
            st.markdown("### Résultat des arbitres")
            for i, a in enumerate(st.session_state["far_arbitres"]):
                nom_complet = f"{a['Prénom']} {a['Nom']}"
                st.markdown(f"**{nom_complet}**")

                col1, col2 = st.columns([3, 3])
                statut = col1.selectbox("Statut", ["Présent", "Absent excusé", "Absent non excusé"], key=f"r_statut_{i}")
                commentaire = col2.text_input("Commentaire", key=f"r_comment_{i}")

                statuts[nom_complet] = statut
                commentaires[nom_complet] = commentaire.strip()


            observations = st.text_area("Observations générales (facultatif)")
            submit = st.form_submit_button("Enregistrer")

            if submit and nom_test:
                for i, a in enumerate(st.session_state["far_arbitres"]):
                    rass = json.loads(a.get("Rassemblements", "") or "[]")
                    rass = [r for r in rass if r.get("Nom") != nom_test]
                    nom_complet = f"{a['Prénom']} {a['Nom']}"
                    rass.append({
                        "Nom": nom_test,
                        "Type": "Test physique",
                        "Date": date_test.strftime("%d/%m/%Y"),
                        "Statut": statuts[nom_complet],
                        "Observations": observations,
                        "Observations individuelles": commentaires.get(f"{a['Prénom']} {a['Nom']}", "")
                    })
                    st.session_state["far_arbitres"][i]["Rassemblements"] = json.dumps(rass)
                save_arbitres(st.session_state["far_arbitres"])
                st.success("Test physique enregistré avec succès.")


elif action == "📊 Récapitulatif des rassemblements":
    st.subheader("📊 Récapitulatif des rassemblements")

    # Chargement des rassemblements
    rassemblements = {}
    for i, arbitre in enumerate(st.session_state["far_arbitres"]):
        try:
            rass_list = json.loads(arbitre.get("Rassemblements", "")) if arbitre.get("Rassemblements") else []
        except:
            rass_list = []

        for r in rass_list:
            nom_rass = r.get("Nom")
            if nom_rass not in rassemblements:
                rassemblements[nom_rass] = {
                    "Type": r.get("Type"),
                    "Date_sort": r.get("Date début", r.get("Date", "")),  # pour tri
                    "Dates": f"{r.get('Date début', r.get('Date', ''))} → {r.get('Date fin', '')}".strip(" →"),
                    "Présences": []
                }
            rassemblements[nom_rass]["Présences"].append((arbitre["Prénom"], arbitre["Nom"], r.get("Statut", "Inconnu")))

    if not rassemblements:
        st.info("Aucun rassemblement enregistré.")
    else:
        # Filtrage par type
        all_types = ["Tous"] + sorted(list(set([r["Type"] for r in rassemblements.values()])))
        selected_type = st.selectbox("Filtrer par type", all_types)

        # Tri des rassemblements
        sorted_rass = sorted(
            rassemblements.items(),
            key=lambda x: datetime.strptime(x[1]["Date_sort"], "%d/%m/%Y") if x[1]["Date_sort"] else datetime.min
        )

        for nom, infos in sorted_rass:
            if selected_type != "Tous" and infos["Type"] != selected_type:
                continue

            statuts_valides = ["présent", "réussi", "échec", "échec partiel"]
            present = sum(1 for p in infos["Présences"] if p[2].lower().strip() in statuts_valides)
            absents_exc = sum(1 for p in infos["Présences"] if p[2].lower().strip() == "absent excusé")
            absents_non = sum(1 for p in infos["Présences"] if p[2].lower().strip() == "absent non excusé")
            total = len(infos["Présences"])


            with st.expander(f"📌 {nom} ({infos['Type']}) — {infos['Dates']}"):
                st.markdown(f"👥 {total} arbitres concernés")
                st.markdown(f"✅ {present} présent(s)")
                st.markdown(f"❌ {absents_exc} absent(s) excusé(s), {absents_non} absent(s) non excusé(s)")

                st.markdown("**Participants :**")
                for prenom, nom_arbitre, statut in infos["Présences"]:
                    st.markdown(f"- {prenom} {nom_arbitre} → *{statut}*")

                if st.button(f"🗑️ Supprimer ce rassemblement", key=f"delete_rass_{nom}"):
                    for i, arbitre in enumerate(st.session_state["far_arbitres"]):
                        rass = json.loads(arbitre.get("Rassemblements", "")) if arbitre.get("Rassemblements") else []
                        rass = [r for r in rass if r.get("Nom") != nom]
                        st.session_state["far_arbitres"][i]["Rassemblements"] = json.dumps(rass)
                    save_arbitres(st.session_state["far_arbitres"])
                    st.success(f"Rassemblement '{nom}' supprimé.")
                    st.rerun()

elif action == "🛑 Ajouter des manquements":
    st.subheader("🛑 Ajouter un manquement à un arbitre")

    # Liste des rassemblements disponibles
    all_rass_names = set()
    for a in st.session_state["far_arbitres"]:
        rass_list = json.loads(a.get("Rassemblements", "") or "[]")
        for r in rass_list:
            all_rass_names.add(r["Nom"])
    rass_names = sorted(list(all_rass_names))

    with st.form("form_ajout_manquement"):
        # Sélection de l’arbitre
        arbitres_dict = {f"{a['Prénom']} {a['Nom']}": i for i, a in enumerate(st.session_state["far_arbitres"])}
        nom_sel = st.selectbox("👤 Sélectionner un arbitre", [""] + list(arbitres_dict.keys()))

        # Type de manquement
        type_manq = st.selectbox("📌 Type de manquement", [
            "",  # Option vide
            "Non-réponse à un Google Form",
            "Absence non excusée",
            "Livrable non rendu",
            "Livrable rendu en retard",
            "Retard",
            "Indisponibilité tardive",
            "Autre"
        ])

        # Date (toujours demandé)
        date = st.date_input("📅 Date du manquement")

        # Initialisation des variables
        detail = ""
        commentaire = ""

        # Champ commentaire (optionnel)
        commentaire = st.text_area("🗒️ Commentaire (optionnel)")

        submit = st.form_submit_button("✅ Ajouter le manquement")

        if submit:
            if not nom_sel:
                st.warning("Veuillez sélectionner un arbitre.")
            elif not type_manq:
                st.warning("Veuillez sélectionner un type de manquement.")
            else:
                i = arbitres_dict[nom_sel]
                try:
                    entry = {
                        "Type": type_manq,
                        "Date": date.strftime("%d/%m/%Y"),
                        "Détail": detail
                    }
                    if commentaire.strip():
                        entry["Commentaire"] = commentaire.strip()

                    raw = st.session_state["far_arbitres"][i].get("Manquements", "")
                    if not isinstance(raw, str):
                        raw = "" if raw is None or (isinstance(raw, float) and math.isnan(raw)) else str(raw)
                    mqs = json.loads(raw or "[]")
                    mqs.append(entry)
                    st.session_state["far_arbitres"][i]["Manquements"] = json.dumps(mqs)
                    save_arbitres(st.session_state["far_arbitres"])
                    st.success(f"Manquement ajouté pour {nom_sel}")
                except Exception as e:
                    st.error(f"Erreur lors de l'enregistrement : {e}")



elif action == "📉 Récapitulatif des manquements":
    st.subheader("📉 Récapitulatif des manquements")

    tous_manquements = []

    for i, a in enumerate(st.session_state["far_arbitres"]):
        try:
            mqs = json.loads(a.get("Manquements", "")) if a.get("Manquements") else []
        except:
            mqs = []

        for m in mqs:
            ligne = {
                "Index": i,
                "Nom complet": f"{a['Prénom']} {a['Nom']}",
                "Type": m.get("Type", "Inconnu"),
                "Détail": m.get("Détail", ""),
                "Date": m.get("Date", "")
            }
            tous_manquements.append(ligne)

    if not tous_manquements:
        st.info("Aucun manquement enregistré.")
    else:
        df_mqs = pd.DataFrame(tous_manquements)

        types_dispo = ["Tous"] + sorted(df_mqs["Type"].unique())
        filtre_type = st.selectbox("Filtrer par type de manquement", types_dispo)

        if filtre_type != "Tous":
            df_mqs = df_mqs[df_mqs["Type"] == filtre_type]

        df_mqs = df_mqs.sort_values(by="Date", ascending=False)

        for idx, row in df_mqs.iterrows():
            with st.expander(f"🚫 {row['Nom complet']} – {row['Type']}"):
                st.markdown(f"📅 **Date :** {row['Date']}")
                if row["Détail"]:
                    st.markdown(f"📝 **Détail :** {row['Détail']}")

                if st.button("🗑️ Supprimer ce manquement", key=f"del_manquement_{idx}"):
                    i = row["Index"]
                    try:
                        mqs = json.loads(st.session_state["far_arbitres"][i].get("Manquements", "")) or []
                        mqs = [m for m in mqs if not (
                            m.get("Type") == row["Type"] and
                            m.get("Date") == row["Date"] and
                            m.get("Détail", "") == row["Détail"]
                        )]
                        st.session_state["far_arbitres"][i]["Manquements"] = json.dumps(mqs)
                        save_arbitres(st.session_state["far_arbitres"])
                        st.success("Manquement supprimé.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erreur : {e}")

elif action == "📝 Saisie des examens":
    st.subheader("📝 Saisie des résultats d'examen")

    with st.form("form_saisie_examen"):
        nom_examen = st.text_input("Nom de l'examen")
        date_examen = st.date_input("Date", value=date.today())
        is_probatoire = st.checkbox("✅ Examen probatoire")

        notes = {}
        statuts = {}

        if not is_probatoire:
            note_max = st.number_input("Note maximale", min_value=1.0, max_value=100.0, value=20.0, step=0.5)

        for i, arbitre in enumerate(st.session_state["far_arbitres"]):
            nom_complet = f"{arbitre['Prénom']} {arbitre['Nom']}"
            st.markdown(f"#### {nom_complet}")

            statut = st.selectbox("Statut", ["Présent", "Absent excusé", "Absent non excusé"], key=f"statut_examen_{i}")
            statuts[i] = statut

            if statut != "Présent":
                notes[i] = {"Absent": True, "Statut": statut}
                st.warning(f"Statut : {statut}")
            else:
                if is_probatoire:
                    qcm = st.number_input("Questionnaire théorique (sur 50)", min_value=0.0, max_value=50.0, step=0.5, key=f"qcm_examen_{i}")
                    video = st.number_input("Test vidéo (sur 30)", min_value=0.0, max_value=30.0, step=0.5, key=f"video_examen_{i}")
                    rapport = st.number_input("Rapport disciplinaire (sur 20)", min_value=0.0, max_value=20.0, step=0.5, key=f"rapport_examen_{i}")
                    total = round(qcm + video + rapport, 2)
                    st.markdown(f"📊 **Total : {total}/100**")
                    notes[i] = {
                        "QCM": qcm,
                        "Vidéo": video,
                        "Rapport": rapport,
                        "Total": total,
                        "Absent": False,
                        "Statut": statut
                    }
                else:
                    note_unique = st.number_input(f"Note (sur {note_max})", min_value=0.0, max_value=note_max, step=0.5, key=f"note_unique_{i}")
                    notes[i] = {
                        "Note": note_unique,
                        "Sur": note_max,
                        "Absent": False,
                        "Statut": statut
                    }

        submit_notes = st.form_submit_button("Enregistrer les notes")

        if submit_notes and nom_examen:

            # Supprimer l'examen existant (nom_examen) pour tous les arbitres avant d'enregistrer
            for i, arbitre in enumerate(st.session_state["far_arbitres"]):
                raw = arbitre.get("Examens", "")
                try:
                    examens = json.loads(raw) if isinstance(raw, str) and raw.strip() else []
                except:
                    examens = []

                new_examens = [e for e in examens if e.get("Nom") != nom_examen]
                st.session_state["far_arbitres"][i]["Examens"] = json.dumps(new_examens)

            for i, arbitre in enumerate(st.session_state["far_arbitres"]):
                raw = arbitre.get("Examens", "")
                try:
                    examens = json.loads(raw) if isinstance(raw, str) and raw.strip() else []
                except:
                    examens = []

                # Supprimer un éventuel doublon
                examens = [e for e in examens if e.get("Nom") != nom_examen]

                data = notes.get(i, {})
                examen_dict = {
                    "Nom": nom_examen,
                    "Type": "Examen",
                    "Date": date_examen.strftime("%d/%m/%Y"),
                    "Statut": data.get("Statut", "Présent"),
                    "Type examen": "Probatoire" if is_probatoire else "Classique"
                }

                if not data.get("Absent", False):
                    if is_probatoire:
                        examen_dict.update({
                            "QCM": data.get("QCM"),
                            "Video": data.get("Vidéo"),
                            "Rapport": data.get("Rapport"),
                            "Note": data.get("Total"),
                            "Sur": 100
                        })
                    else:
                        examen_dict.update({
                            "Note": data.get("Note"),
                            "Sur": data.get("Sur")
                        })

                examens.append(examen_dict)
                st.session_state["far_arbitres"][i]["Examens"] = json.dumps(examens)

            save_arbitres(st.session_state["far_arbitres"])
            st.success("✅ Résultats d'examen enregistrés avec succès.")


elif action == "📊 Récapitulatif des examens":
    st.subheader("📊 Récapitulatif des examens")

    import matplotlib.pyplot as plt
    from collections import defaultdict

    examens = []
    for i, a in enumerate(st.session_state["far_arbitres"]):
        try:
            exam_list = json.loads(a.get("Examens", "") or "[]")
        except:
            exam_list = []

        for r in exam_list:
            examens.append({
                "Index": i,
                "Nom complet": f"{a['Prénom']} {a['Nom']}",
                "Catégorie": a.get("Catégorie", ""),
                "Type": r.get("Type examen", "Classique"),
                "Date": r.get("Date"),
                "Nom": r.get("Nom"),
                "Statut": r.get("Statut"),
                "Note": float(r.get("Note", 0)),
                "Sur": float(r.get("Sur", 100)),
                "QCM": r.get("QCM", None),
                "Video": r.get("Video", None),
                "Rapport": r.get("Rapport", None),
            })




    if not examens:
        st.info("Aucun examen enregistré.")
        st.stop()

    df_examens = pd.DataFrame(examens)

    # Statistiques globales
    st.markdown("### 📈 Statistiques globales")

    nb_exams_uniques = df_examens["Nom"].nunique()
    nb_absents = df_examens[df_examens["Statut"] != "Présent"].groupby("Nom").ngroups
    taux_presence = 100 * (nb_exams_uniques - nb_absents) / nb_exams_uniques if nb_exams_uniques else 0

    st.markdown(f"- Nombre total d'examens : **{nb_exams_uniques}**")
    st.markdown(f"- Nombre d'examens avec au moins un absent : **{nb_absents}**")
    st.markdown(f"- Taux de présence globale : **{taux_presence:.1f}%**")


    # Graphiques
    st.markdown("### 📊 Graphiques")

    # Moyenne par arbitre
    df_notes = df_examens[df_examens["Statut"] == "Présent"].copy()
    df_notes["Pourcentage"] = df_notes["Note"] / df_notes["Sur"] * 100

    df_examens["Pourcentage"] = 100 * df_examens["Note"] / df_examens["Sur"]
    moyennes = df_notes.groupby("Nom complet")["Pourcentage"].mean().sort_values()
    totaux = df_examens.groupby("Nom complet")["Pourcentage"].sum().sort_values()

    fig1, ax1 = plt.subplots(figsize=(8, 4))
    moyennes.plot(kind="barh", ax=ax1)
    ax1.set_title("Moyenne % par arbitre (présents uniquement)")
    st.pyplot(fig1)

    fig2, ax2 = plt.subplots(figsize=(8, 4))
    totaux.plot(kind="barh", ax=ax2)
    ax2.set_title("Total % par arbitre (absences pénalisantes)")
    st.pyplot(fig2)

    # Classement par catégorie
    st.markdown("### 🏆 Classement par catégorie")

    cat_options = df_examens["Catégorie"].dropna().unique().tolist()
    selected_cat = st.selectbox("Filtrer par catégorie", ["Toutes"] + sorted(cat_options))

    if selected_cat != "Toutes":
        df_notes = df_notes[df_notes["Catégorie"] == selected_cat]
        df_examens = df_examens[df_examens["Catégorie"] == selected_cat]

    classement_moyenne = df_notes.groupby("Nom complet")["Pourcentage"].mean().sort_values(ascending=False)
    classement_total = df_examens.groupby("Nom complet")["Pourcentage"].sum().sort_values(ascending=False)

    st.markdown("#### Classement par moyenne")
    for i, (nom, score) in enumerate(classement_moyenne.items(), 1):
        st.markdown(f"{i}. **{nom}** – {score:.1f}%")

    st.markdown("#### Classement par total")
    for i, (nom, score) in enumerate(classement_total.items(), 1):
        st.markdown(f"{i}. **{nom}** – {score:.1f} points cumulés")

    # Table complète
    st.markdown("### 📋 Détails des examens")
    # 🎯 Filtrer par nom d'examen
    examens_disponibles = df_examens["Nom"].unique().tolist()
    selected_exam = st.selectbox("📝 Filtrer par examen", ["Tous"] + sorted(examens_disponibles))

    df_display = df_examens.copy()
    df_display["Note (%)"] = df_display["Pourcentage"].map(lambda x: f"{x:.1f}%")

    # 🧠 Sélection des colonnes selon type d’examen
    colonnes_communes = ["Nom", "Date", "Nom complet", "Catégorie", "Type", "Statut"]

    if selected_exam != "Tous":
        df_display = df_display[df_display["Nom"] == selected_exam]

    # Teste si examen probatoire
    is_probatoire = (
        df_display["Type"].iloc[0] == "Probatoire"
        if not df_display.empty
        else False
    )

    if is_probatoire:
        colonnes = colonnes_communes + [c for c in ["QCM", "Video", "Rapport", "Note", "Sur", "Note (%)"] if c in df_display.columns]
    else:
        colonnes = colonnes_communes + [c for c in ["Note", "Sur", "Note (%)"] if c in df_display.columns]

    # Filtrage des colonnes disponibles
    df_display = df_display[[col for col in colonnes if col in df_display.columns]]

    st.dataframe(df_display, use_container_width=True)

    if selected_exam != "Tous":
        df_exam_unique = df_examens[df_examens["Nom"] == selected_exam].copy()

        st.markdown(f"## 📄 Résultats détaillés pour l'examen : {selected_exam}")

        # 🔢 Classement
        classement = df_exam_unique[df_exam_unique["Statut"] == "Présent"].copy()
        classement["Pourcentage"] = 100 * classement["Note"] / classement["Sur"]
        classement = classement.sort_values(by="Pourcentage", ascending=False)

        st.markdown("### 🏅 Classement des arbitres (présents uniquement)")
        for rang, (_, row) in enumerate(classement.iterrows(), start=1):
            st.markdown(f"{rang}. **{row['Nom complet']}** – {row['Note']:.1f} / {row['Sur']}")


        # 📊 Graphique
        import matplotlib.pyplot as plt

        st.markdown("### 📊 Graphique des résultats")
        fig, ax = plt.subplots(figsize=(10, 4))
        classement.set_index("Nom complet")["Pourcentage"].plot(kind="barh", ax=ax)
        ax.invert_yaxis()
        ax.set_xlabel("Note en %")
        ax.set_title(f"Résultats à l'examen « {selected_exam} »")
        st.pyplot(fig)

        # 🧾 Tableau détaillé exportable
        st.markdown("### 📋 Tableau des résultats")
        colonnes_affichage = ["Nom complet", "Catégorie", "Statut", "Type", "Note", "Sur", "Pourcentage"]

        if df_exam_unique["Type"].iloc[0] == "Probatoire":
            colonnes_affichage = ["Nom complet", "Catégorie", "Statut", "QCM", "Video", "Rapport", "Note", "Sur", "Pourcentage"]

        df_export = df_exam_unique.copy()
        df_export["Pourcentage"] = df_export["Pourcentage"].map(lambda x: round(x, 1))
        df_export = df_export.reset_index(drop=True)  # pour éviter décalages d'index
        df_export.insert(0, "Rang", df_export.index + 1)


        st.dataframe(df_export[colonnes_affichage], use_container_width=True)

        # 📥 Bouton de téléchargement
        from io import BytesIO

        buffer = BytesIO()
        df_export[colonnes_affichage].to_excel(buffer, index=False)
        buffer.seek(0)

        st.download_button(
            label="📥 Télécharger les résultats",
            data=buffer,
            file_name=f"Résultats_{selected_exam}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    st.markdown("### 🧹 Supprimer un examen")

    noms_examens_dispo = df_examens["Nom"].unique().tolist()
    examen_a_supprimer = st.selectbox("📌 Sélectionnez un examen à supprimer", [""] + noms_examens_dispo)

    if examen_a_supprimer:
        if st.button("🗑️ Supprimer cet examen pour tous les arbitres"):
            nb_suppr = 0
            for i, arbitre in enumerate(st.session_state["far_arbitres"]):
                examens = json.loads(arbitre.get("Examens", "") or "[]")
                new_examens = [e for e in examens if e.get("Nom") != examen_a_supprimer]
                if len(new_examens) != len(examens):  # s'il y a eu une suppression
                    st.session_state["far_arbitres"][i]["Examens"] = json.dumps(new_examens)
                    nb_suppr += 1

            if nb_suppr > 0:
                save_arbitres(st.session_state["far_arbitres"])
                st.success(f"✅ Examen '{examen_a_supprimer}' supprimé pour {nb_suppr} arbitre(s).")
                st.rerun()
            else:
                st.info("Aucun arbitre n'était concerné par cet examen.")

if action == "📎 Déposer un rapport d'observation":
    st.subheader("📎 Déposer un rapport pour un arbitre")

    # Sélection de l’arbitre
    arbitres_dict = {f"{a['Prénom']} {a['Nom']}": i for i, a in enumerate(st.session_state["far_arbitres"])}

    with st.form("upload_rapport_form"):
        nom_sel = st.selectbox("👤 Sélectionner un arbitre", [""] + list(arbitres_dict.keys()))
        uploaded_file = st.file_uploader("📄 Déposer un rapport (PDF, Word...)", type=["pdf", "docx", "doc"])

        submit = st.form_submit_button("✅ Enregistrer le rapport")

    if submit and nom_sel and uploaded_file:
        temp_path = "temp_upload.pdf"
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        assert os.path.exists(temp_path), "Le fichier PDF n’a pas été généré"

        try:
            folder_id = st.secrets["GOOGLE_DRIVE_FOLDER_ID"]
            url_partage = upload_to_drive(temp_path, uploaded_file.name, parent_folder_id=folder_id)
            os.remove(temp_path)
        except Exception as e:
            st.error(f"Erreur lors de l'upload sur Google Drive : {e}")
            st.stop()

        rapports = json.loads(st.session_state["far_arbitres"][arbitres_dict[nom_sel]].get("Rapports", "[]"))
        rapports.append({
            "nom_original": uploaded_file.name,
            "url": url_partage
        })
        st.session_state["far_arbitres"][arbitres_dict[nom_sel]]["Rapports"] = json.dumps(rapports, ensure_ascii=False)

        save_arbitres(st.session_state["far_arbitres"])
        st.success("✅ Rapport envoyé sur Google Drive.")
        st.rerun()





elif action == "👤 Fiche arbitre":
    st.subheader("👤 Rechercher un arbitre")
    noms = [f"{a['Prénom']} {a['Nom']}" for a in st.session_state["far_arbitres"]]
    sel = st.selectbox("Sélectionnez un arbitre", [""] + sorted(noms))
    
    if sel:
        a = next(a for a in st.session_state["far_arbitres"] if f"{a['Prénom']} {a['Nom']}" == sel)

        st.markdown(f"### {sel} — {a.get('Catégorie')} — Âge : {a.get('Âge')}")
        st.markdown(f"- 📅 Date de naissance : {a.get('Date de naissance')}")
        st.markdown(f"- 🏠 Club : {a.get('Club')}")
        tel = str(a.get("Téléphone", "")).strip()
        if tel and not tel.startswith("0") and len(tel) == 9:
            tel = "0" + tel
        email = a.get("Email", "—")

        st.markdown(f"- 📞 Tel : {tel}  |  ✉️ Email : {email}")

        # === Rassemblements ===
        rass = json.loads(a.get("Rassemblements", "") or "[]")
        if rass:
            st.markdown("### 📋 Rassemblements")

            # Convertir en DataFrame
            df_rass = pd.DataFrame(rass)

            # Fusionner Date début et Date fin en une date d’affichage
            df_rass["Date"] = df_rass.apply(
                lambda row: row["Date"] if "Date" in row and pd.notna(row["Date"])
                else row.get("Date début", ""), axis=1
            )

           

            # Convertir en datetime pour tri
            df_rass["Date_dt"] = pd.to_datetime(df_rass["Date"], format="%d/%m/%Y", errors="coerce")


            # Regrouper par type
            for type_rass in sorted(df_rass["Type"].dropna().unique()):
                st.markdown(f"#### 🗂️ {type_rass}")

                df_sous = df_rass[df_rass["Type"] == type_rass].copy()
                df_sous = df_sous.sort_values("Date_dt")

                # Affichage propre
                df_sous = df_sous.rename(columns={"Observations individuelles": "Observation"})
                for col in ["Nom", "Date", "Type", "Statut", "Observation"]:
                    if col not in df_sous.columns:
                        df_sous[col] = ""

                df_sous_affiche = df_sous[["Nom", "Date", "Type", "Statut", "Observation"]].reset_index(drop=True)
                df_sous_affiche.index = df_sous_affiche.index + 1
                st.dataframe(df_sous_affiche, use_container_width=True)


        # === Examens ===
            exam = json.loads(a.get("Examens", "") or "[]")
            if exam:
                st.markdown("#### 🧠 Examens")
                for e in exam:
                    if "Note" not in e and "Total" in e:
                        e["Note"] = e["Total"]
                    if "Type examen" not in e:
                        e["Type examen"] = ""
                    if e.get("Type examen") != "Probatoire":
                        e["QCM"] = e["Video"] = e["Rapport"] = ""

                df_exam = pd.DataFrame(exam)
                colonnes = ["Nom", "Date", "Type examen", "Statut", "Note", "Sur", "QCM", "Video", "Rapport"]
                df_exam = df_exam[[c for c in colonnes if c in df_exam.columns]]
                df_exam.index = df_exam.index + 1
                st.dataframe(df_exam, use_container_width=True)


                # === Position globale dans la catégorie ===
                nom_complet = f"{a['Prénom']} {a['Nom']}"
                cat = a.get("Catégorie")
                total_points = 0

                # Total de l'arbitre courant
                exam = json.loads(a.get("Examens", "") or "[]")
                for e in exam:
                    try:
                        total_points += float(e.get("Note", e.get("Total", 0)))
                    except:
                        continue

                # Calcul du classement global
                classement = []
                for autre in st.session_state["far_arbitres"]:
                    if autre.get("Catégorie") != cat:
                        continue
                    try:
                        raw = autre.get("Examens", "")
                        autres_exams = json.loads(raw if isinstance(raw, str) else "")
                        total = sum(float(e.get("Note", e.get("Total", 0))) for e in autres_exams)
                        nom_autre = f"{autre.get('Prénom')} {autre.get('Nom')}"
                        classement.append((nom_autre, total))
                    except:
                        continue

                classement.sort(key=lambda x: x[1], reverse=True)
                rang = next((i + 1 for i, (n, _) in enumerate(classement) if n == nom_complet), None)
                total_cat = len(classement)

                if rang:
                    st.markdown(f"### 📊 Position globale dans la catégorie **{cat}** : {rang}ᵉ sur {total_cat}")



        # === Manquements ===
        val = a.get("Manquements", "")
        try:
            mqs = json.loads(val if isinstance(val, str) else "")
        except:
            mqs = []

        if mqs:
            st.markdown("#### 🚫 Manquements")
            df_manq = pd.DataFrame(mqs)
            colonnes = ["Date", "Type", "Détail", "Commentaire"]
            df_manq = df_manq[[c for c in colonnes if c in df_manq.columns]]
            df_manq.index = df_manq.index + 1
            st.dataframe(df_manq, use_container_width=True)
        

        st.markdown("### 📎 Rapports associés")
        try:
            arbitre_id = f"{a['Nom'].upper()}_{a['Prénom']}".replace(" ", "_")
            rapports = list_rapports_for_arbitre(arbitre_id)
            if rapports:
                for nom, url in rapports:
                    st.markdown(f"- [{nom}]({url})", unsafe_allow_html=True)
                    col1, col2 = st.columns([3, 1])
                    with col2:
                        if st.button("🗑️ Supprimer", key=f"sup_{nom}"):
                            filepath = url.split("/object/public/rapports/")[-1]
                            delete_rapport_from_supabase(filepath)
                            st.success(f"Rapport « {nom} » supprimé.")
                            st.experimental_rerun()
            else:
                st.info("Aucun rapport n’est encore associé à cet arbitre.")
        except Exception as e:
            st.error(f"Erreur lors de la récupération des rapports : {e}")


        # === Boutons Word ===
        from docx import Document
        from io import BytesIO
        from docx.shared import Pt
        from datetime import datetime
        import json

        def create_doc_for_arbitre(a, arbitres_liste):
            doc = Document()
            nom_complet = f"{a['Prénom']} {a['Nom']}"

            # === En-tête ===
            doc.add_heading(nom_complet, level=1)
            doc.add_paragraph(f"Catégorie : {a.get('Catégorie', '—')}")
            doc.add_paragraph(f"Âge : {a.get('Âge', '—')}")
            doc.add_paragraph(f"Né(e) le : {a.get('Date de naissance', '—')}")
            doc.add_paragraph(f"Club : {a.get('Club', '—')}")
            doc.add_paragraph(f"Téléphone : {str(a.get('Téléphone', '—'))}")
            doc.add_paragraph(f"Email : {a.get('Email', '—')}")

            # === Examens ===
            exams_raw = a.get("Examens", "")
            try:
                exams = json.loads(exams_raw if isinstance(exams_raw, str) else "")
            except:
                exams = []

            if exams:
                doc.add_heading("Examens", level=2)
                table = doc.add_table(rows=1, cols=8)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "Nom"
                hdr[1].text = "Date"
                hdr[2].text = "Type"
                hdr[3].text = "Statut"
                hdr[4].text = "Note"
                hdr[5].text = "QCM"
                hdr[6].text = "Vidéo"
                hdr[7].text = "Rapport"

                total_points = 0
                for e in exams:
                    note = float(e.get("Note", e.get("Total", 0)))
                    total_points += note

                    row = table.add_row().cells
                    row[0].text = e.get("Nom", "?")
                    row[1].text = e.get("Date", "?")
                    row[2].text = e.get("Type examen", "")
                    row[3].text = e.get("Statut", "")
                    row[4].text = f"{note} / {e.get('Sur', 100)}"
                    row[5].text = str(e.get("QCM", "")) if e.get("Type examen") == "Probatoire" else ""
                    row[6].text = str(e.get("Video", "")) if e.get("Type examen") == "Probatoire" else ""
                    row[7].text = str(e.get("Rapport", "")) if e.get("Type examen") == "Probatoire" else ""

            else:
                total_points = 0

            # === Rassemblements ===
            rass_raw = a.get("Rassemblements", "")
            try:
                rass = json.loads(rass_raw if isinstance(rass_raw, str) else "")
            except:
                rass = []

            if rass:
                rass.sort(key=lambda r: datetime.strptime(r.get("Date", "01/01/1900"), "%d/%m/%Y"))
                doc.add_heading("Rassemblements", level=2)
                table = doc.add_table(rows=1, cols=5)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "Nom"
                hdr[1].text = "Date"
                hdr[2].text = "Type"
                hdr[3].text = "Statut"
                hdr[4].text = "Observation individuelle"

                for r in rass:
                    row = table.add_row().cells
                    row[0].text = r.get("Nom", "?")
                    row[1].text = r.get("Date", "?")
                    row[2].text = r.get("Type", "")
                    row[3].text = r.get("Statut", "")
                    row[4].text = r.get("Observations individuelles", "")

            # === Manquements ===
            val = a.get("Manquements", "")
            try:
                mqs = json.loads(val if isinstance(val, str) else "")
            except:
                mqs = []

            if mqs:
                doc.add_heading("Manquements", level=2)
                table = doc.add_table(rows=1, cols=3)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "Date"
                hdr[1].text = "Type"
                hdr[2].text = "Détail"

                for m in mqs:
                    row = table.add_row().cells
                    row[0].text = m.get("Date", "")
                    row[1].text = m.get("Type", "")
                    row[2].text = m.get("Détail", "")

            # === Position globale dans la catégorie ===
            cat = a.get("Catégorie")
            if cat and total_points > 0:
                classement = []
                for autre in arbitres_liste:
                    if autre.get("Catégorie") != cat:
                        continue
                    try:
                        raw = autre.get("Examens", "")
                        autres_exams = json.loads(raw if isinstance(raw, str) else "")
                        total = sum(float(e.get("Note", e.get("Total", 0))) for e in autres_exams)
                        nom_autre = f"{autre.get('Prénom')} {autre.get('Nom')}"
                        classement.append((nom_autre, total))
                    except:
                        continue

                classement.sort(key=lambda x: x[1], reverse=True)
                rang = next((i+1 for i, (n, _) in enumerate(classement) if n == nom_complet), None)
                total_cat = len(classement)

                if rang:
                    doc.add_paragraph(f"📊 Position globale dans la catégorie **{cat}** : {rang}ᵉ sur {total_cat}")

            return doc

        st.subheader("📎 Rapport d'observation")

        uploaded_file = st.file_uploader("Ajouter un rapport PDF", type=["pdf"], key=f"rapport_{a['Nom']}_{a['Prénom']}")

        if uploaded_file:
            try:
                arbitre_id = f"{a['Nom'].upper()}_{a['Prénom']}".replace(" ", "_")
                url = upload_rapport_to_supabase(uploaded_file, arbitre_id)
                st.success("Rapport téléversé avec succès !")
                st.markdown(f"[📄 Voir le rapport]({url})")

                # Mise à jour dans la fiche (temporaire en session)
                a["Rapport observation"] = url

            except Exception as e:
                st.error(f"Erreur : {e}")

        if a.get("Rapport observation"):
            st.markdown(f"📁 Rapport déjà enregistré : [Voir le rapport]({a['Rapport observation']})")

        st.divider()
        col1, col2 = st.columns(2)
        with col1:
            if st.button("📄 Générer fiche Word"):
                doc = create_doc_for_arbitre(a, st.session_state["far_arbitres"])
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    "Télécharger fiche Word",
                    data=buffer,
                    file_name=f"{sel}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        with col2:
            if st.button("📁 Générer fiches Word (tous les arbitres)"):
                doc_all = Document()
                for arbitre in st.session_state["far_arbitres"]:
                    fiche = create_doc_for_arbitre(arbitre, st.session_state["far_arbitres"])
                    for p in fiche.paragraphs:
                        doc_all.add_paragraph(p.text)
                    doc_all.add_page_break()

                buffer = BytesIO()
                doc_all.save(buffer)
                buffer.seek(0)
                st.download_button(
                    "Télécharger fiches (tous)",
                    data=buffer,
                    file_name="fiches_arbitres.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )






# === AFFICHAGE LISTE ===
st.markdown("---")
st.subheader("📋 Liste des arbitres")
for a in st.session_state["far_arbitres"]:
    st.markdown(f"- **{a['Prénom']} {a['Nom']}** | {a['Catégorie']} | Âge : {a.get('Âge', 'N/A')}")

